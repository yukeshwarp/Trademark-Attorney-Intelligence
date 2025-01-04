# Version - 6.1  (Code Updated of Checking for Non-matching Class number == GPT 4o mini)

from fileinput import filename
import time
import os
import streamlit as st 
import pandas as pd
import fitz  # PyMuPDF
from pydantic import BaseModel, Field, ValidationError
from typing import List, Dict, Union
import base64
from docx import Document  
from docx.shared import Pt
from io import BytesIO
import re, ast
from dotenv import load_dotenv
load_dotenv()  

class TrademarkDetails(BaseModel):
    trademark_name: str = Field(description="The name of the Trademark", example="DISCOVER")
    status: str = Field(description="The Status of the Trademark", example="Registered")
    serial_number: str = Field(description="The Serial Number of the trademark from Chronology section", example="87−693,628")
    international_class_number: List[int] = Field(description="The International class number or Nice Classes number of the trademark from Goods/Services section or Nice Classes section", example=[18])
    owner: str = Field(description="The owner of the trademark", example="WALMART STORES INC")
    goods_services: str = Field(description="The goods/services from the document", example="LUGGAGE AND CARRYING BAGS; SUITCASES, TRUNKS, TRAVELLING BAGS, SLING BAGS FOR CARRYING INFANTS, SCHOOL BAGS; PURSES; WALLETS; RETAIL AND ONLINE RETAIL SERVICES")
    page_number: int = Field(description="The page number where the trademark details are found in the document", example=3)
    registration_number: Union[str, None] = Field(description="The Registration number of the trademark from Chronology section", example="5,809,957")
    design_phrase: str = Field(description="The design phrase of the trademark", example="THE MARK CONSISTS OF THE STYLIZED WORD 'MINI' FOLLOWED BY 'BY MOTHERHOOD.'", default="")
    

def preprocess_text(text: str) -> str:
    text = re.sub(r'\s+', ' ', text).strip()
    text = re.sub(r'[\u2013\u2014]', '-', text)
    return text

def is_correct_format_code1(page_text: str) -> bool:
    required_fields = ["Status:", "Goods/Services:"] # , "Last Reported Owner:"
    return all(field in page_text for field in required_fields)

def is_correct_format_code2(page_text: str) -> bool:
    required_fields = ["Register", "Nice Classes", "Goods & Services"]
    return all(field in page_text for field in required_fields)

def extract_trademark_details_code1(document_chunk: str) -> Dict[str, Union[str, List[int]]]:
    try:
        from openai import AzureOpenAI
        azure_endpoint = os.getenv("AZURE_ENDPOINT")  
        api_key = os.getenv("AZURE_API_KEY")  
          
        client = AzureOpenAI(  
            azure_endpoint=azure_endpoint,  
            api_key=api_key,   
            api_version="2024-10-01-preview",
        )  
         
        # messages = [  
        #     {  
        #         "role": "system",  
        #         "content": """You are a highly skilled data extraction specialist with expertise in parsing complex trademark research reports. Your primary goal is to accurately extract specified information from provided documents and present it in a structured format, strictly adhering to the given instructions and formatting requirements.  
                
        #         Please read the **Trademark Research Report** provided below and extract the following details for **each trademark** listed in the report. Return the results as a formatted string for each trademark, containing the following properties:  
        
        #         - **Trademark Name**: The name of the trademark.  
        #         - **Status**: The status of the trademark (e.g., "Registered", "Pending", "Abandoned") as indicated in the report.  
        #         - **Serial Number**: The serial number of the trademark.  
        #         - **International Class Number**: A list of integers representing the International Class numbers associated with the trademark.  
        #         - **Owner**: The owner of the trademark.  
        #         - **Goods & Services**: The description of the goods/services associated with the trademark.  
        #         - **Filed Date**: The filing date of the trademark in the format "MMM DD, YYYY".  
        #         - **Registration Number**: The registration number of the trademark, if available. If not available, state "Not available in the document."  
                
        #         **Instructions:**  
        #         - **Iterate through each trademark entry** in the report, starting from the **USPTO Summary** section.  
        #         - Extract the required details for each trademark, referring to both the summary and the detailed entries that follow.  
        #         - Use the exact property names provided above in your output.  
        #         - For fields that are missing or not available, state "Not available in the document."  
        #         - Ensure that:  
        #         - Dates are formatted as "MMM DD, YYYY". If the date is spelled out (e.g., "November 27, 2024"), convert it accordingly to "Nov 27, 2024".  
        #         - Numbers are represented appropriately (e.g., integers for class numbers and registration numbers).  
        #         - **International Class Number** is a single integer or list of integers.  
        #         - For owners with multiple parties, list all parties in a single string separated by semicolons.  
        
        #         **Example Output:**  
        #         ```  
        #         Trademark Name: SLIK Global Filings Status: PENDING FILED AS USE APPLICATION Serial Number: 98-602,112 International Class Number: 3 Owner: SLIK DE VENEZUELA C.A. VENEZUELA CORPORATION Goods & Services: Cosmetics; hair gel; hair wax; hair styling gel; non-medicated cosmetics Filed Date: JUN 14, 2024 Registration Number: Not available in the document.  
        #         ```  
        #         """  
        #     },  
        #     {  
        #         "role": "user",  
        #         "content": f"""  **Document:** {document_chunk}"""  
        #     }  
        # ]  
        
        messages=[  
                    {"role": "system", "content": "You are a data extraction specialist proficient in parsing trademark documents."},  
                    {"role": "user", "content": f"""  
                        Extract the following details from the provided trademark document and present them in the exact format specified:  

                        - Trademark Name  
                        - Status  
                        - Serial Number  
                        - International Class Number (as a list of integers)  
                        - Owner  
                        - Goods & Services  
                        - Filed Date (format: MMM DD, YYYY, e.g., Jun 14, 2024)  
                        - Registration Number  

                        **Instructions:**  
                        - Return the results in the following format, replacing the example data with the extracted information:
                        - Ensure the output matches this format precisely.  
                        - Do not include any additional text or explanations.  

                        **Document to extract from:**  
                        {document_chunk}  
                        """}  
                        ]  
        
        response = client.chat.completions.create(  
                model="gpt-4o-mini",  
                messages=messages,  
                temperature=0,  
                max_tokens=4000,  
        )  
        extracted_text = response.choices[0].message.content
        
        # if extracted_text and extracted_text != "[]":
        #     st.write(extracted_text)
            
        details = {}
        for line in extracted_text.split("\n"):
            if ":" in line:
                key, value = line.split(":", 1)
                details[key.strip().lower().replace(" ", "_")] = value.strip()
        return details
    
    except Exception as e:
        print(f"An error occurred: {e}")

    
def extract_registration_number(document: str) -> str:
    """ Extract the registration number from the Chronology section """
    match = re.search(r'Chronology:.*?Registration Number:\s*([\d,]+)', document, re.DOTALL)
    if match:
        return match.group(1).strip()
    return "No registration number presented in document"

def extract_trademark_details_code2(page_text: str) -> Dict[str, Union[str, List[int]]]:
    details = {}

    trademark_name_match = re.search(r"\d+\s*/\s*\d+\s*\n\s*\n\s*([A-Za-z0-9'&!,\-. ]+)\s*\n", page_text)
    if trademark_name_match:
        details["trademark_name"] = trademark_name_match.group(1).strip()
    else:
        trademark_name_match = re.search(r"(?<=\n)([A-Za-z0-9'&!,\-. ]+)(?=\n)", page_text)
        details["trademark_name"] = trademark_name_match.group(1).strip() if trademark_name_match else ""

    status_match = re.search(r'Status\s*(?:\n|:\s*)([A-Za-z]+)', page_text, re.IGNORECASE)
    details["status"] = status_match.group(1).strip() if status_match else ""

    owner_match = re.search(r'Holder\s*(?:\n|:\s*)(.*)', page_text, re.IGNORECASE)
    if owner_match:
        details["owner"] = owner_match.group(1).strip()
    else:
        owner_match = re.search(r'Owner\s*(?:\n|:\s*)(.*)', page_text, re.IGNORECASE)
        details["owner"] = owner_match.group(1).strip() if owner_match else ""
        
        

    nice_classes_match = re.search(r'Nice Classes\s*[\s:]*\n((?:\d+(?:,\s*\d+)*)\b)', page_text, re.IGNORECASE)
    if nice_classes_match:
        nice_classes_text = nice_classes_match.group(1)
        nice_classes = [int(cls.strip()) for cls in nice_classes_text.split(",")]
        details["international_class_number"] = nice_classes
    else:
        details["international_class_number"] = []
        


    serial_number_match = re.search(r'Application#\s*(.*)', page_text, re.IGNORECASE)
    details["serial_number"] = serial_number_match.group(1).strip() if serial_number_match else ""

    goods_services_match = re.search(r'Goods & Services\s*(.*?)(?=\s*G&S translation|$)', page_text, re.IGNORECASE | re.DOTALL)
    details["goods_services"] = goods_services_match.group(1).strip() if goods_services_match else ""
    
    registration_number_match = re.search(r'Registration#\s*(.*)', page_text, re.IGNORECASE)
    details["registration_number"] = registration_number_match.group(1).strip() if registration_number_match else ""
    
    # Description
    design_phrase = re.search(r'Description\s*(.*?)(?=\s*Applicant|Owner|Holder|$)', page_text, re.IGNORECASE | re.DOTALL)
    details["design_phrase"] = design_phrase.group(1).strip() if design_phrase else "No Design phrase presented in document"
    

    return details

def read_pdf(file_path: str, exclude_header_footer: bool = True) -> str:
    document_text = ""
    with fitz.open(file_path) as pdf_document:
        for page_num in range(pdf_document.page_count):
            page = pdf_document.load_page(page_num)
            if exclude_header_footer:
                rect = page.rect
                x0 = rect.x0
                y0 = rect.y0 + rect.height * 0.1
                x1 = rect.x1
                y1 = rect.y1 - rect.height * 0.1
                page_text = page.get_text("text", clip=(x0, y0, x1, y1))
            else:
                page_text = page.get_text()
            document_text += page_text
    return document_text

def split_text(text: str, max_tokens: int = 1500) -> List[str]:
    chunks = []
    current_chunk = []
    current_length = 0

    for line in text.split('\n'):
        line_length = len(line.split())
        if current_length + line_length > max_tokens:
            chunks.append('\n'.join(current_chunk))
            current_chunk = [line]
            current_length = line_length
        else:
            current_chunk.append(line)
            current_length += line_length

    if current_chunk:
        chunks.append('\n'.join(current_chunk))

    return chunks

def parse_international_class_numbers(class_numbers: str) -> List[int]:
    numbers = class_numbers.split(',')
    return [int(num.strip()) for num in numbers if num.strip().isdigit()]

# def extract_international_class_numbers_and_goods_services(document: str) -> Dict[str, Union[List[int], str]]:
#     """ Extract the International Class Numbers and Goods/Services from the document """
#     class_numbers = []
#     goods_services = []
#     pattern = r'International Class (\d+): (.*?)(?=\nInternational Class \d+:|\n[A-Z][a-z]+:|\nLast Reported Owner:|\Z)'
#     matches = re.findall(pattern, document, re.DOTALL)
#     for match in matches:
#         class_number = int(match[0])
#         class_numbers.append(class_number)
#         goods_services.append(f"Class {class_number}: {match[1].strip()}")
#     return {
#         "international_class_numbers": class_numbers,
#         "goods_services": "\n".join(goods_services)
#     }

def extract_international_class_numbers_and_goods_services(document: str, start_page: int, pdf_document: fitz.Document) -> Dict[str, Union[List[int], str]]:
    """ Extract the International Class Numbers and Goods/Services from the document over a range of pages """
    class_numbers = []
    goods_services = []
    combined_text = ""

    for i in range(start_page, min(start_page + 6, pdf_document.page_count)):
        page = pdf_document.load_page(i)
        page_text = page.get_text()
        combined_text += page_text
        if "Last Reported Owner:" in page_text:
            break

    pattern = r'International Class (\d+): (.*?)(?=\nInternational Class \d+:|\n[A-Z][a-z]+:|\nLast Reported Owner:|Disclaimers:|\Z)'
    matches = re.findall(pattern, combined_text, re.DOTALL)
    for match in matches:
        class_number = int(match[0])
        class_numbers.append(class_number)
        goods_services.append(f"Class {class_number}: {match[1].strip()}")

    if "sexual" in goods_services or "sex" in goods_services:
        goods_services = replace_disallowed_words(goods_services)

    return {
        "international_class_numbers": class_numbers,
        "goods_services": "\n".join(goods_services)
    }

def extract_design_phrase(document: str, start_page: int, pdf_document: fitz.Document) -> Dict[str, Union[List[int], str]]:
    """ Extract the design phrase from the document """
    combined_texts = ""
    for i in range(start_page, min(start_page + 8, pdf_document.page_count)):
        page = pdf_document.load_page(i)
        page_text = page.get_text()
        combined_texts += page_text
        if "Filing Correspondent:" in page_text:
            break
        
    pattern = r'Design Phrase:\s*(.*?)(?=Other U\.S\. Registrations:|Filing Correspondent:|Group:|USPTO Page:|$)'
    match = re.search(pattern, document, re.DOTALL) 
    if match:
        design_phrase = match.group(1).strip()
        # Remove any newline characters within the design phrase
        design_phrase = ' '.join(design_phrase.split())
        return design_phrase
    return "No Design phrase presented in document"

    
def parse_trademark_details(document_path: str) -> List[Dict[str, Union[str, List[int]]]]:
    with fitz.open(document_path) as pdf_document:
        all_extracted_data = []
        for page_num in range(pdf_document.page_count):
            page = pdf_document.load_page(page_num)
            page_text = page.get_text()
            
            if is_correct_format_code1(page_text):
                preprocessed_chunk = preprocess_text(page_text)
                extracted_data = extract_trademark_details_code1(preprocessed_chunk)
                additional_data = extract_international_class_numbers_and_goods_services(page_text, page_num, pdf_document)
                registration_number = extract_registration_number(page_text)
                design_phrase = extract_design_phrase(page_text, page_num, pdf_document)
                
                if extracted_data:
                    extracted_data["page_number"] = page_num + 1
                    extracted_data.update(additional_data)
                    extracted_data["design_phrase"] = design_phrase
                    all_extracted_data.append(extracted_data)
                    extracted_data["registration_number"] = registration_number
                    
                trademark_list = []
                for i, data in enumerate(all_extracted_data, start=1):
                    try:
                        trademark_name = data.get("trademark_name", "").split(',')[0].strip()
                        if "Global Filings" in trademark_name:
                            trademark_name = trademark_name.split("Global Filings")[0].strip()
                        owner = data.get("owner", "").split(',')[0].strip()
                        status = data.get("status", "").split(',')[0].strip()
                        serial_number = data.get("serial_number", "")
                        international_class_number = data.get("international_class_numbers", [])
                        goods_services = data.get("goods_services", "")
                        page_number = data.get("page_number", "")
                        registration_number = data.get("registration_number", "No registration number presented in document")
                        design_phrase = data.get("design_phrase", "No Design phrase presented in document")

                        # If crucial fields are missing, attempt to re-extract the values
                        if not trademark_name or not owner or not status or not international_class_number:
                            preprocessed_chunk = preprocess_text(data.get("raw_text", ""))
                            extracted_data = extract_trademark_details_code1(preprocessed_chunk)
                            trademark_name = extracted_data.get("trademark_name", trademark_name).split(',')[0].strip()
                            if "Global Filings" in trademark_name:
                                trademark_name = trademark_name.split("Global Filings")[0].strip()
                            owner = extracted_data.get("owner", owner).split(',')[0].strip()
                            status = extracted_data.get("status", status).split(',')[0].strip()
                            international_class_number = parse_international_class_numbers(extracted_data.get("international_class_number", "")) or international_class_number
                            registration_number = extracted_data.get("registration_number", registration_number).split(',')[0].strip()

                        trademark_details = TrademarkDetails(
                            trademark_name=trademark_name,
                            owner=owner,
                            status=status,
                            serial_number=serial_number,
                            international_class_number=international_class_number,
                            goods_services=goods_services,
                            page_number=page_number,
                            registration_number=registration_number,
                            design_phrase=design_phrase
                        )                        
                        trademark_info = {
                            "trademark_name": trademark_details.trademark_name,
                            "owner": trademark_details.owner,
                            "status": trademark_details.status,
                            "serial_number": trademark_details.serial_number,
                            "international_class_number": trademark_details.international_class_number,
                            "goods_services": trademark_details.goods_services,
                            "page_number": trademark_details.page_number,
                            "registration_number":trademark_details.registration_number,
                            "design_phrase": trademark_details.design_phrase
                        }
                        print(trademark_info)
                        print("_____________________________________________________________________________________________________________________________")
                        trademark_list.append(trademark_info)
                    except ValidationError as e:
                        print(f"Validation error for trademark {i}: {e}")
                                    
            else :
                if not is_correct_format_code2(page_text):
                    continue

                extracted_data = extract_trademark_details_code2(page_text)
                if extracted_data:
                    extracted_data["page_number"] = page_num + 1
                    all_extracted_data.append(extracted_data)

                trademark_list = []
                for i, data in enumerate(all_extracted_data, start=1):
                    try:
                        trademark_details = TrademarkDetails(
                            trademark_name=data.get("trademark_name", ""),
                            owner=data.get("owner", ""),
                            status=data.get("status", ""),
                            serial_number=data.get("serial_number", ""),
                            international_class_number=data.get("international_class_number", []),
                            goods_services=data.get("goods_services", ""),
                            page_number=data.get("page_number", 0),
                            registration_number=data.get("registration_number", ""),
                            design_phrase=data.get("design_phrase", "")
                        )
                        if (trademark_details.trademark_name != "" and trademark_details.owner != "" and trademark_details.status != "" and trademark_details.goods_services != ""):
                                trademark_info = {
                                    "trademark_name": trademark_details.trademark_name,
                                    "owner": trademark_details.owner,
                                    "status": trademark_details.status,
                                    "serial_number": trademark_details.serial_number,
                                    "international_class_number": trademark_details.international_class_number,
                                    "goods_services": trademark_details.goods_services,
                                    "page_number": trademark_details.page_number,
                                    "registration_number":trademark_details.registration_number,
                                    "design_phrase":trademark_details.design_phrase,
                                }
                                
                                trademark_list.append(trademark_info)
                    except ValidationError as e:
                        print(f"Validation error for trademark {i}: {e}")

        return trademark_list

from azure.core.exceptions import HttpResponseError 

# total_tokens_per_trademark = prompt_tokens_per_trademark + response_tokens_per_trademark  
# max_batch_size = (max_token_limit - buffer_tokens) // total_tokens_per_trademark  
# If len(existing_trademarks) is less than max_batch_size, set batch_size to len(existing_trademarks).
# If len(existing_trademarks) is larger, use max_batch_size
    
def compare_trademarks(existing_trademark: List[Dict[str, Union[str, List[int]]]], proposed_name: str, proposed_class: str, proposed_goods_services: str) -> List[Dict[str, int]]:
    proposed_classes = [int(c.strip()) for c in proposed_class.split(',')]
    results = []  

    # Set the batch size based on token limits  
    batch_size = 5  # Adjust this number as needed
    
    for i in range(0, len(existing_trademarks), batch_size):  
        batch = existing_trademarks[i:i + batch_size] 
        messages=[
                {"role": "system", "content": """You are an experienced Trademark Attorney specializing in trademark law and intellectual property rights. You hold a Juris Doctor (J.D.) degree and have over 10 years of experience in conducting comprehensive trademark searches, analyzing potential conflicts, and advising clients on trademark strategies. Your expertise includes assessing the likelihood of confusion between trademarks, understanding trademark classifications, and providing detailed legal opinions.
                                                Role: As a Trademark Attorney, your task is to analyze existing trademark and proposed trademark.

                                                Condition 1: Trademark Name Comparison
                                                - Condition 1A: Are the existing trademark name and the proposed trademark name in conflict with respect to Distinctiveness, Strength of the Marks, and Similarity in Appearance, Sound, and Meaning?

                                                If the existing trademark in the user-provided input satisfies:
                                                - Special Case: If the existing trademark status is "Cancelled" or "Abandoned," it will automatically be considered a conflict grade of "Low," but you should still provide reasoning for any potential conflicts.
                                                - If the existing trademark satisfies Condition 1A, then the conflict grade should be "Name-Match."
                                                - If the existing trademark does not satisfy Condition 1A, then the conflict grade should be "Low."

                                                Format of the Response:
                                                Reasoning for Conflict: Provide reasoning for the conflict in bullet points. In your reasoning, if the goods, services, and industries are exactly the same, list the overlaps. You should determine whether the goods/services overlap, including classes (whether they are the same as the proposed trademark or not). Consider whether the trademark names are identical (character-for-character matches), phonetic equivalents, if the name is in the primary position (first word in the phrase), or if it is not in the primary position of the existing trademark. If it is not in the primary position, it is not conflicting. Also, consider standard plural forms for subject goods and whether the goods may be related or not. Reasoning should be based on the provided information. Do not provide any hypothetical reasoning.
                                                Note: Also mention if the existing trademark and the proposed trademark are not in the same Class number in the Reasoning for Conflict.
                                                
                                                Step 0: Identifying Potential Conflicts
                                                - What is the existing trademark?
                                                - What is the status of the existing trademark? 
                                                - What is the proposed trademark?

                                                Step 1: Check the Status of the Existing Trademark:
                                                - If the existing trademark is "Cancelled" or "Expired" or "Abandoned," assign the conflict grade as "Low." And skip other conditions.

                                                Step 2: **Trademark Name Comparison:**
                                                - Evaluate if there is a conflict between the existing trademark name and the proposed trademark name based on the following:
                                                    - **Distinctiveness and Strength of the Marks:** Are the trademarks distinctive or similar in strength?
                                                    - **Similarity in Appearance, Sound, and Meaning:** Do the names look, sound, or mean the same?

                                                Step 3: **Consider Special Cases and Additional Factors:**
                                                - Consider standard plural forms for the subject goods and whether goods may be related or not.
                                                - If there is no similarity in name, class, or overlapping goods/services, assign the conflict grade "Low."

                                                Format of the Response:
                                                - **Reasoning for Conflict:** Provide reasoning in bullet points. Base your reasoning only on the provided information. Clearly mention if the existing and proposed trademarks are not in the same Class number.

                                                Note:
                                                - Do not provide any hypothetical reasoning. The conflict grade should be based solely on the facts given.
                                                
                                                Example Analysis Using the Steps : 
                                                - Trademark Name: Unlock Brisk's Bold Flavors
                                                - Trademark Status: REGISTERED
                                                - Proposed Trademark: Unlock Hidden Flavors

                                                Reasoning for Conflict:
                                                Step 1: Status Check: 
                                                - The existing trademark status is not "Cancelled" or "Expired" or "Abandoned,". Proceeding to name comparison.

                                                Step 2: Trademark Name Comparison:
                                                - Both trademarks share the distinctive word "Unlock" in the primary position, which creates a similarity in appearance, sound, and meaning.
                                                - The phrase "Unlock Flavors" forms the core part of both trademarks, creating a strong conceptual similarity.
                                                - While the words "Brisk's Bold" and "Hidden" differ, they serve as modifiers to the common term "Flavors." The similarity is substantial because the focus of both trademarks is on the concept of "Unlocking Flavors," which could confuse consumers regarding the source or affiliation of the products/services.

                                                Step 3: No special cases apply (such as status being "Cancelled" or "Abandoned").
                                                
                                                Conflict Reason:
                                                Reasoning for Conflict:
                                                The shared use of the distinctive phrase "Unlock Flavors" as the primary conceptual focus in both trademarks creates a strong similarity in appearance, sound, and meaning. This overlap is significant enough to potentially confuse consumers about the source or affiliation of the goods or services, thereby assigning the conflict grade as "Name-Match."
                                                
                                                - Conflict Grade: Name-Match
                                                """
                                                },            
                {"role": "user", "content": """Compare the following existing and proposed trademarks and determine the conflict grade.\n
                                                Existing Trademark:\n
                                                Name: SCOOPT'D\n
                                                Status: Registered\n
                                                                                            
                                                Proposed Trademark:\n
                                                Name: SCOOP-A-PALOOZA\n """
                },
                {"role": "assistant", "content":""" 
    Reasoning for Conflict:
    Step 1: Status Check:
    - The existing trademark status is not "Cancelled," "Expired," or "Abandoned." Proceeding to name comparison.

    Step 2: Trademark Name Comparison:
    - Both trademarks share the distinctive word "SCOOP," which creates a similarity in appearance, sound, and meaning.
    - The word "SCOOP" is the dominant and distinctive part of both trademarks, leading to a conceptual similarity.
    - The existing trademark "SCOOPT'D" and the proposed trademark "SCOOP-A-PALOOZA" both emphasize the idea of "Scoop," likely related to ice cream or a similar product, which could confuse consumers regarding the source or affiliation.
    - The additional elements ("T'D" and "A-PALOOZA") differ but serve as suffixes or modifiers to the common term "SCOOP."

    Step 3: No special cases apply (such as status being "Cancelled" or "Abandoned").

    Conflict Reason:
    Reasoning for Conflict:
    The shared use of the distinctive word "SCOOP" as the core focus in both trademarks creates a strong similarity in appearance, sound, and meaning. This overlap is significant enough to potentially confuse consumers about the source or affiliation of the goods or services, thereby assigning the conflict grade as "Name-Match."

    - Conflict Grade: Name-Match
    """
                },
                {"role": "user", "content": f"""Compare the following existing and proposed trademarks and determine the conflict grade.\n
                                                Existing Trademark:\n
                                                Name: {existing_trademark['trademark_name']}\n
                                                Status: {existing_trademark['status']}\n
                                                
                                                Proposed Trademark:\n
                                                Name: {proposed_name}\n """
                }
            ]
        
        from openai import AzureOpenAI
        azure_endpoint = os.getenv("AZURE_ENDPOINT")  
        api_key = os.getenv("AZURE_API_KEY")  
            
        client = AzureOpenAI(  
            azure_endpoint=azure_endpoint,  
            api_key=api_key,   
            api_version="2024-10-01-preview",
        )
        
        try: 
            response_reasoning = client.chat.completions.create(  
                        model="gpt-4o-mini",  
                        messages=messages,  
                        temperature=0,  
                        max_tokens=2000,   
                        top_p = 1
                        )

            Treasoning = response_reasoning.choices[0].message.content
            print(Treasoning)
            print("_____________________________________________________________________________________________________________________________")
            if Treasoning is not None:
                # if ("Reasoning for Conflict:" in Treasoning or "Based on the analysis:" in Treasoning) and "Conflict Grade:" in Treasoning:
                #     if "Reasoning for Conflict:" in Treasoning:
                #         reasoning = Treasoning.split("Reasoning for Conflict:", 1)[1].strip()
                #     else:
                #         reasoning = Treasoning.split("Based on the analysis:", 1)[1].strip()

                if "Conflict Grade:" in Treasoning:
                    reasoning = Treasoning
                    conflict_grade = Treasoning.split("Conflict Grade:", 1)[1].strip() 
                    progress_bar.progress(60)
                
                    return {
                        'Trademark name': existing_trademark['trademark_name'],
                        'Trademark status': existing_trademark['status'],
                        'Trademark owner': existing_trademark['owner'],
                        'Trademark class Number': existing_trademark['international_class_number'],
                        'Trademark serial number' : existing_trademark['serial_number'],
                        'Trademark registration number' : existing_trademark['registration_number'],
                        'Trademark design phrase' : existing_trademark['design_phrase'],
                        'conflict_grade': conflict_grade,
                        'reasoning': reasoning
                    }
                else:
                    print("Check Response")  
                    st.error(f"Trademark: {existing_trademark['trademark_name']}, \n Output from LLM: {Treasoning}")  
                    return {
                        'Trademark name': existing_trademark['trademark_name'],
                        'Trademark status': existing_trademark['status'],
                        'Trademark owner': existing_trademark['owner'],
                        'Trademark class Number': existing_trademark['international_class_number'],
                        'Trademark serial number' : existing_trademark['serial_number'],
                        'Trademark registration number' : existing_trademark['registration_number'],
                        'Trademark design phrase' : existing_trademark['design_phrase'],
                        'conflict_grade': "Not Defined",
                        'reasoning': "Not Defined"
                    }
            else:
                print("Check Response")  
                st.info(f"Trademark: {existing_trademark['trademark_name']}, \n Output from LLM: {Treasoning}")  
                return {
                    'Trademark name': existing_trademark['trademark_name'],
                    'Trademark status': existing_trademark['status'],
                    'Trademark owner': existing_trademark['owner'],
                    'Trademark class Number': existing_trademark['international_class_number'],
                    'Trademark serial number' : existing_trademark['serial_number'],
                    'Trademark registration number' : existing_trademark['registration_number'],
                    'Trademark design phrase' : existing_trademark['design_phrase'],
                    'conflict_grade': "Not Defined",
                    'reasoning': "Not Defined"
                }
        except HttpResponseError as e:  
            print(f"HTTP error occurred: {e.message}")
            st.warning(f"Trademark: {existing_trademark['trademark_name']}")  
        except Exception as e:  
            print(f"An unexpected error occurred: {str(e)}") 
            st.warning(f"Trademark: {existing_trademark['trademark_name']}")  

def replace_disallowed_words(text):
    disallowed_words = {
        "sexual": "xxxxxx",
        "sex": "xxx",
    }
    for word, replacement in disallowed_words.items():
        text = text.replace(word, replacement)
    # Ensure single paragraph output
    text = " ".join(text.split())
    return text

def compare_trademarks2(existing_trademark: List[Dict[str, Union[str, List[int]]]], proposed_name: str, proposed_class: str, proposed_goods_services: str) -> List[Dict[str, int]]:
    proposed_classes = [int(c.strip()) for c in proposed_class.split(',')]
    messages=[
            {"role": "system", "content": """You are a trademark attorney tasked with determining a conflict grade based on the given conditions. You should assign a conflict grade of "Name-Match" or "Low" to the existing trademark and respond with only "Name-Match", or "Low".
                                            Conditions for Determining Conflict Grades:

                                            Condition 1: Trademark Name Comparison
                                            - Condition 1A: Are the existing trademark name and the proposed trademark name in conflict with respect to Distinctiveness, Strength of the Marks, and Similarity in Appearance, Sound, and Meaning?

                                            If the existing trademark in the user-provided input satisfies:
                                            - Special Case: If the existing trademark status is "Cancelled" or "Abandoned," it will automatically be considered a conflict grade of "Low," but you should still provide reasoning for any potential conflicts.
                                            - If the existing trademark satisfies Condition 1A, then the conflict grade should be "Name-Match."
                                            - If the existing trademark does not satisfy Condition 1A, then the conflict grade should be "Low."

                                            Format of the Response:
                                            Reasoning for Conflict: Provide reasoning for the conflict in bullet points. In your reasoning, if the goods, services, and industries are exactly the same, list the overlaps. You should determine whether the goods/services overlap, including classes (whether they are the same as the proposed trademark or not). Consider whether the trademark names are identical (character-for-character matches), phonetic equivalents, if the name is in the primary position (first word in the phrase), or if it is not in the primary position of the existing trademark. If it is not in the primary position, it is not conflicting. Also, consider standard plural forms for subject goods and whether the goods may be related or not. Reasoning should be based on the provided information. Do not provide any hypothetical reasoning.
                                            Note: Also mention if the existing trademark and the proposed trademark are not in the same Class number in the Reasoning for Conflict.
                                            
                                            Step 0: Identifying Potential Conflicts
                                            - What is the existing trademark?
                                            - What is the status of the existing trademark? 
                                            - What is the proposed trademark?

                                            Step 1: Check the Status of the Existing Trademark:
                                            - If the existing trademark is "Cancelled" or "Expired" or "Abandoned," assign the conflict grade as "Low." And skip other conditions.

                                            Step 2: **Trademark Name Comparison:**
                                            - Evaluate if there is a conflict between the existing trademark name and the proposed trademark name based on the following:
                                                - **Distinctiveness and Strength of the Marks:** Are the trademarks distinctive or similar in strength?
                                                - **Similarity in Appearance, Sound, and Meaning:** Do the names look, sound, or mean the same?

                                            Step 3: **Consider Special Cases and Additional Factors:**
                                            - Consider standard plural forms for the subject goods and whether goods may be related or not.
                                            - If there is no similarity in name, class, or overlapping goods/services, assign the conflict grade "Low."

                                            Format of the Response:
                                            - **Reasoning for Conflict:** Provide reasoning in bullet points. Base your reasoning only on the provided information. Clearly mention if the existing and proposed trademarks are not in the same Class number.

                                            Note:
                                            - Do not provide any hypothetical reasoning. The conflict grade should be based solely on the facts given.
                                            
                                            Example Analysis Using the Steps : 
                                            - Trademark Name: Unlock Brisk's Bold Flavors
                                            - Trademark Status: REGISTERED
                                            - Proposed Trademark: Unlock Hidden Flavors

                                            Reasoning for Conflict:
                                            Step 1: Status Check: 
                                            - The existing trademark status is not "Cancelled" or "Expired" or "Abandoned,". Proceeding to name comparison.

                                            Step 2: Trademark Name Comparison:
                                            - Both trademarks share the distinctive word "Unlock" in the primary position, which creates a similarity in appearance, sound, and meaning.
                                            - The phrase "Unlock Flavors" forms the core part of both trademarks, creating a strong conceptual similarity.
                                            - While the words "Brisk's Bold" and "Hidden" differ, they serve as modifiers to the common term "Flavors." The similarity is substantial because the focus of both trademarks is on the concept of "Unlocking Flavors," which could confuse consumers regarding the source or affiliation of the products/services.

                                            Step 3: No special cases apply (such as status being "Cancelled" or "Abandoned").
                                            
                                            Conflict Reason:
                                            Reasoning for Conflict:
                                            The shared use of the distinctive phrase "Unlock Flavors" as the primary conceptual focus in both trademarks creates a strong similarity in appearance, sound, and meaning. This overlap is significant enough to potentially confuse consumers about the source or affiliation of the goods or services, thereby assigning the conflict grade as "Name-Match."
                                            
                                            - Conflict Grade: Name-Match
                                            """
                                            },            
            {"role": "user", "content": """Compare the following existing and proposed trademarks and determine the conflict grade.\n
                                            Existing Trademark:\n
                                            Name: SCOOPT'D\n
                                            Status: Registered\n
                                                                                        
                                            Proposed Trademark:\n
                                            Name: SCOOP-A-PALOOZA\n """
            },
            {"role": "assistant", "content":""" 
Reasoning for Conflict:
Step 1: Status Check:
- The existing trademark status is not "Cancelled," "Expired," or "Abandoned." Proceeding to name comparison.

Step 2: Trademark Name Comparison:
- Both trademarks share the distinctive word "SCOOP," which creates a similarity in appearance, sound, and meaning.
- The word "SCOOP" is the dominant and distinctive part of both trademarks, leading to a conceptual similarity.
- The existing trademark "SCOOPT'D" and the proposed trademark "SCOOP-A-PALOOZA" both emphasize the idea of "Scoop," likely related to ice cream or a similar product, which could confuse consumers regarding the source or affiliation.
- The additional elements ("T'D" and "A-PALOOZA") differ but serve as suffixes or modifiers to the common term "SCOOP."

Step 3: No special cases apply (such as status being "Cancelled" or "Abandoned").

Conflict Reason:
Reasoning for Conflict:
The shared use of the distinctive word "SCOOP" as the core focus in both trademarks creates a strong similarity in appearance, sound, and meaning. This overlap is significant enough to potentially confuse consumers about the source or affiliation of the goods or services, thereby assigning the conflict grade as "Name-Match."

- Conflict Grade: Name-Match
"""
            },
            {"role": "user", "content": f"""Compare the following existing and proposed trademarks and determine the conflict grade.\n
                                            Existing Trademark:\n
                                            Name: {existing_trademark['trademark_name']}\n
                                            Status: {existing_trademark['status']}\n
                                            
                                            Proposed Trademark:\n
                                            Name: {proposed_name}\n """
            }
        ]

    from openai import AzureOpenAI
    azure_endpoint = os.getenv("AZURE_ENDPOINT")  
    api_key = os.getenv("AZURE_API_KEY")  
        
    client = AzureOpenAI(  
        azure_endpoint=azure_endpoint,  
        api_key=api_key,   
        api_version="2024-10-01-preview",
    )
                
    response_reasoning = client.chat.completions.create(  
                        model="gpt-4o-mini",  
                        messages=messages,  
                        temperature=0,  
                        max_tokens=4095,  
                        top_p = 1
                    )

    Treasoning = response_reasoning.choices[0].message.content
    print(Treasoning)
    print("_____________________________________________________________________________________________________________________________")
    reasoning = Treasoning.split("Reasoning for Conflict:", 1)[1].strip()
    conflict_grade = Treasoning.split("Conflict Grade:", 1)[1].strip() 
    progress_bar.progress(70)
    
    return {
        'Trademark name': existing_trademark['trademark_name'],
        'Trademark status': existing_trademark['status'],
        'Trademark owner': existing_trademark['owner'],
        'Trademark class Number': existing_trademark['international_class_number'],
        'Trademark serial number' : existing_trademark['serial_number'],
        'Trademark registration number' : existing_trademark['registration_number'],
        'Trademark design phrase' : existing_trademark['design_phrase'],
        'conflict_grade': conflict_grade,
        'reasoning': reasoning
    }
    

def extract_proposed_trademark_details(file_path: str) -> Dict[str, Union[str, List[int]]]:
    """ Extract proposed trademark details from the given input format """
    proposed_details = {}
    with fitz.open(file_path) as pdf_document:
        if pdf_document.page_count > 0:
            page = pdf_document.load_page(0)
            page_text = preprocess_text(page.get_text())
            
    name_match = re.search(r'Mark Searched:\s*(.*?)(?=\s*Client Name:)', page_text, re.IGNORECASE | re.DOTALL)
    if name_match:
        proposed_details["proposed_trademark_name"] = name_match.group(1).strip()

    goods_services_match = re.search(r'Goods/Services:\s*(.*?)(?=\s*Trademark Research Report)', page_text, re.IGNORECASE | re.DOTALL)
    if goods_services_match:
        proposed_details["proposed_goods_services"] = goods_services_match.group(1).strip()
    
    # Use LLM to find the international class number based on goods & services
    if "proposed_goods_services" in proposed_details:
        goods_services = proposed_details["proposed_goods_services"]
        class_numbers = find_class_numbers(goods_services)
        proposed_details["proposed_nice_classes_number"] = class_numbers
    
    return proposed_details

def find_class_numbers(goods_services: str) -> List[int]:
    """ Use LLM to find the international class numbers based on goods & services """
        # Initialize AzureChatOpenAI
    
    from openai import AzureOpenAI
    azure_endpoint = os.getenv("AZURE_ENDPOINT")  
    api_key = os.getenv("AZURE_API_KEY")  
          
    client = AzureOpenAI(  
        azure_endpoint=azure_endpoint,  
        api_key=api_key,   
        api_version="2024-10-01-preview",
    )

    # messages = [  
    # {  
    #     "role": "system",  
    #     "content": "You are an expert in classifying goods and services into their respective International Class Numbers according to the Nice Classification."  
    # },  
    # {  
    #     "role": "user",  
    #     "content": f"""Given the following goods/services, identify all applicable International Class Numbers.  
    #     **Instructions:**  

    #     - Provide **only** a comma-separated list of class numbers (e.g., `3`, `18,35`).  
    #     - Do not include any additional text, explanations, or formatting.  
    #     - Ensure the class numbers are integers and separated by commas without spaces.  

    #     **Examples:**  

    #     - Goods/Services: "SKIN CARE PREPARATIONS; COSMETICS; BABY CARE PRODUCTS..."  
    #     - Response: 3  

    #     - Goods/Services: "LUGGAGE AND CARRYING BAGS; SUITCASES, TRUNKS..."  
    #     - Response: 18,35  

    #     **Goods/Services to Classify:**  
    #     {goods_services}  
    #     """  
    #         }  
    #     ]  
    
    messages = [  
                {  
                    "role": "system",  
                    "content": "You are an expert in classifying goods and services into their respective International Class Numbers according to the Nice Classification."  
                },  
                {  
                    "role": "user",  
                    "content": f"""  
                                Identify all applicable International Class Numbers for the following goods/services. Think through the classification process step by step.  

                                **Instructions:**  
                                - **Step 1**: Analyze each item in the goods/services list.  
                                - **Step 2**: Determine the appropriate class number(s) for each item.  
                                - **Step 3**: Compile a list of unique class numbers.  
                                - **Step 4**: Provide **only** a comma-separated list of two-digit class numbers (e.g., `03`, `18,35`).  
                                - Do not include any additional text, explanations, or formatting.  

                                **Goods/Services:**  
                                {goods_services}  
                                """  
                }  
            ]  
    
    response = client.chat.completions.create(  
                        model="gpt-4o-mini",  
                        messages=messages,  
                        temperature=0,  
                        max_tokens=150,  
    )  
    
    class_numbers_str = response.choices[0].message.content
    
    # Extracting class numbers and removing duplicates
    class_numbers = re.findall(r'(?<!\d)\d{2}(?!\d)', class_numbers_str)  # Look for two-digit numbers
    class_numbers = ','.join(set(class_numbers))  # Convert to set to remove duplicates, then join into a single string
    
    return class_numbers

def extract_proposed_trademark_details2(file_path: str) -> Dict[str, Union[str, List[int]]]:
    """ Extract proposed trademark details from the first page of the document """
    proposed_details = {}
    with fitz.open(file_path) as pdf_document:
        if pdf_document.page_count > 0:
            page = pdf_document.load_page(0)
            page_text = preprocess_text(page.get_text())
            
            name_match = re.search(r'Name:\s*(.*?)(?=\s*Nice Classes:)', page_text)
            if name_match:
                proposed_details["proposed_trademark_name"] = name_match.group(1).strip()
                
            nice_classes_match = re.search(r'Nice Classes:\s*(\d+(?:,\s*\d+)*)', page_text)
            if nice_classes_match:
                proposed_details["proposed_nice_classes_number"] = nice_classes_match.group(1).strip()
            
            goods_services_match = re.search(r'Goods & Services:\s*(.*?)(?=\s*Registers|$)', page_text, re.IGNORECASE | re.DOTALL)
            if goods_services_match:
                proposed_details["proposed_goods_services"] = goods_services_match.group(1).strip()
    
    return proposed_details

def list_conversion(proposed_class: str) -> List[int]:
    
    from openai import AzureOpenAI
    azure_endpoint = os.getenv("AZURE_ENDPOINT")  
    api_key = os.getenv("AZURE_API_KEY")  
        
    client = AzureOpenAI(  
        azure_endpoint=azure_endpoint,  
        api_key=api_key,   
        api_version="2024-10-01-preview",
    )

    messages = [  
                {  
                    "role": "system",  
                    "content": "You are a helpful assistant that converts strings of class numbers into a Python list of integers."  
                },  
                {  
                    "role": "user",  
                    "content": f"""  
                            Convert the following string of class numbers into a Python list of integers. Follow the steps to ensure accuracy.  

                            **Instructions:**  
                            - **Step 1**: Split the input string by commas to separate each number.  
                            - **Step 2**: Remove any whitespace and convert each number to an integer.  
                            - **Step 3**: Combine the integers into a Python list.  
                            - **Step 4**: Respond **only** with the Python list of integers (e.g., `[15, 89]`).  
                            - Do not include any additional text or commentary.  

                            **Input:**  
                            "{proposed_class}"  
                            """  
                }  
            ]  
    
    response = client.chat.completions.create(  
                        model="gpt-4o-mini",  
                        messages=messages,  
                        temperature=0.2,  
                        max_tokens=150,  
    )  

    lst_class = response.choices[0].message.content
    class_value = ast.literal_eval(lst_class)
            
    return class_value

# Streamlit App  
st.title("Trademark Document Parser Version 6.2")  
  
# File upload  
uploaded_files = st.sidebar.file_uploader("Choose PDF files", type="pdf", accept_multiple_files=True)  
  
if uploaded_files:  
    if st.sidebar.button("Check Conflicts", key="check_conflicts"):  
        total_files = len(uploaded_files)  
        progress_bar = st.progress(0)  
        # progress_label.text(f"Progress: 0%")  --- Needed to set

        for i, uploaded_file in enumerate(uploaded_files):  
            # Save uploaded file to a temporary file path  
            temp_file_path = f"temp_{uploaded_file.name}"  
            with open(temp_file_path, "wb") as f:  
                f.write(uploaded_file.read())  
                
            start_time = time.time()
            
            sp = True
            proposed_trademark_details = extract_proposed_trademark_details(temp_file_path)  
                            
            if proposed_trademark_details:  
                proposed_name = proposed_trademark_details.get('proposed_trademark_name', 'N')  
                proposed_class = proposed_trademark_details.get('proposed_nice_classes_number')  
                proposed_goods_services = proposed_trademark_details.get('proposed_goods_services', 'N') 
                if (proposed_goods_services != 'N'): 
                    with st.expander(f"Proposed Trademark Details for {uploaded_file.name}"):  
                            st.write(f"Proposed Trademark name: {proposed_name}")  
                            st.write(f"Proposed class-number: {proposed_class}")  
                            st.write(f"Proposed Goods & Services: {proposed_goods_services}") 
                    class_list = list_conversion(proposed_class) 
                else :
                    st.write("______________________________________________________________________________________________________________________________")
                    st.write(f"Sorry, unable to generate report due to insufficient information about goods & services in the original trademark report : {uploaded_file.name}")
                    st.write("______________________________________________________________________________________________________________________________")
                    sp = False
            else:  
                
                proposed_trademark_details = extract_proposed_trademark_details2(temp_file_path)  
                
                if proposed_trademark_details:  
                    proposed_name = proposed_trademark_details.get('proposed_trademark_name', 'N')  
                    proposed_class = proposed_trademark_details.get('proposed_nice_classes_number')  
                    proposed_goods_services = proposed_trademark_details.get('proposed_goods_services', 'N')  
                    if (proposed_goods_services != 'N'): 
                        with st.expander(f"Proposed Trademark Details for {uploaded_file.name}"):  
                                st.write(f"Proposed Trademark name: {proposed_name}")  
                                st.write(f"Proposed class-number: {proposed_class}")  
                                st.write(f"Proposed Goods & Services: {proposed_goods_services}") 
                        class_list = list_conversion(proposed_class)  
                    else :
                        st.write("______________________________________________________________________________________________________________________________")
                        st.write(f"Sorry, unable to generate report due to insufficient information about goods & services in the original trademark report : {uploaded_file.name}")
                        st.write("______________________________________________________________________________________________________________________________")
                        sp = False
                else :  
                    st.error(f"Unable to extract Proposed Trademark Details for {uploaded_file.name}") 
                    sp = False 
                    continue  
            
            if (sp):    
                for i in range(1,21):
                    progress_bar.progress(i)
                    
                progress_bar.progress(25)
                # Initialize AzureChatOpenAI
                
                # s_time = time.time()
                
                existing_trademarks = parse_trademark_details(temp_file_path)
                for i in range(25,46):
                    progress_bar.progress(i)  
                    
                progress_bar.progress(50)
                st.success(f"Existing Trademarks Data Extracted Successfully for {uploaded_file.name}!")  
                
                # e_time = time.time()
                # elap_time = e_time - s_time
                # elap_time = elap_time // 60 
                # st.write(f"Time taken for extraction: {elap_time} mins")

                # e_time = time.time()  
                # elap_time = e_time - s_time  
                # st.write(f"Time taken: {elap_time} seconds")  
                
                # Display extracted details              
                
                nfiltered_list = []
                unsame_class_list = []
                
                # Iterate over each JSON element in trademark_name_list  
                for json_element in existing_trademarks:  
                    class_numbers = json_element["international_class_number"]  
                # Check if any of the class numbers are in class_list  
                    if any(number in class_list for number in class_numbers):  
                        nfiltered_list.append(json_element)
                    else:
                        unsame_class_list.append(json_element)
                    
                existing_trademarks = nfiltered_list  
                existing_trademarks_unsame =  unsame_class_list
                     
                high_conflicts = []
                moderate_conflicts = []
                low_conflicts = []
                Name_Matchs = []
                
                lt = len(existing_trademarks)
                
                for existing_trademark in existing_trademarks:  
                    conflict = compare_trademarks(existing_trademark, proposed_name, proposed_class, proposed_goods_services)  
                    if conflict is not None:
                        if conflict['conflict_grade'] == "High":  
                            high_conflicts.append(conflict)  
                        elif conflict['conflict_grade'] == "Moderate":  
                            moderate_conflicts.append(conflict)  
                        else:  
                            low_conflicts.append(conflict)  
                        
                for existing_trademarks in existing_trademarks_unsame:
                    conflict = compare_trademarks2(existing_trademarks, proposed_name, proposed_class, proposed_goods_services)  

                    if conflict['conflict_grade'] == "Name-Match":  
                        Name_Matchs.append(conflict)  
                    else:  
                        print("Low")
                        # low_conflicts.append(conflict) 
                        
                st.sidebar.write("_________________________________________________")
                st.sidebar.subheader("\n\nConflict Grades : \n")  
                st.sidebar.markdown(f"File: {proposed_name}")  
                st.sidebar.markdown(f"Total number of conflicts: {len(high_conflicts) + len(moderate_conflicts) + len(Name_Matchs) + len(low_conflicts)}")
                st.sidebar.markdown(f"High Conflicts: {len(high_conflicts)}")  
                st.sidebar.markdown(f"Moderate Conflicts: {len(moderate_conflicts)}")  
                st.sidebar.markdown(f"Name Match's Conflicts: {len(Name_Matchs)}")  
                st.sidebar.markdown(f"Low Conflicts: {len(low_conflicts)}")  
                st.sidebar.write("_________________________________________________")
    
                document = Document()  
                
                document.add_heading(f'Trademark Conflict List for {proposed_name} (VERSION - 6.2) :')            
                document.add_paragraph(f"\n\nTotal number of conflicts: {len(high_conflicts) + len(moderate_conflicts) + len(Name_Matchs) + len(low_conflicts)}\n- High Conflicts: {len(high_conflicts)}\n- Moderate Conflicts: {len(moderate_conflicts)}\n- Name Match's Conflicts: {len(Name_Matchs)}\n- Low Conflicts: {len(low_conflicts)}\n")  
                
                if len(high_conflicts) > 0:  
                            document.add_heading('Trademarks with High Conflicts:', level=2)  
                            # Create a pandas DataFrame from the JSON list    
                            df_high = pd.DataFrame(high_conflicts) 
                            df_high = df_high.drop(columns=['Trademark serial number','Trademark registration number','Trademark design phrase','reasoning'])  
                            # Create a table in the Word document    
                            table_high = document.add_table(df_high.shape[0] + 1, df_high.shape[1])
                            # Set a predefined table style (with borders)  
                            table_high.style = 'TableGrid'  # This is a built-in style that includes borders  
                            # Add the column names to the table    
                            for i, column_name in enumerate(df_high.columns):  
                                table_high.cell(0, i).text = column_name  
                            # Add the data to the table    
                            for i, row in df_high.iterrows():  
                                for j, value in enumerate(row):  
                                    table_high.cell(i + 1, j).text = str(value)

                if len(moderate_conflicts) > 0:  
                            document.add_heading('Trademarks with Moderate Conflicts:', level=2)  
                            # Create a pandas DataFrame from the JSON list    
                            df_moderate = pd.DataFrame(moderate_conflicts)
                            df_moderate = df_moderate.drop(columns=['Trademark serial number','Trademark registration number','Trademark design phrase','reasoning'])  
                            # Create a table in the Word document    
                            table_moderate = document.add_table(df_moderate.shape[0] + 1, df_moderate.shape[1])
                            # Set a predefined table style (with borders)  
                            table_moderate.style = 'TableGrid'  # This is a built-in style that includes borders  
                            # Add the column names to the table    
                            for i, column_name in enumerate(df_moderate.columns):  
                                table_moderate.cell(0, i).text = column_name  
                            # Add the data to the table    
                            for i, row in df_moderate.iterrows():  
                                for j, value in enumerate(row):  
                                    table_moderate.cell(i + 1, j).text = str(value)
                                    
                if len(Name_Matchs) > 0:  
                            document.add_heading("Trademarks with Name Match's Conflicts:", level=2)  
                            # Create a pandas DataFrame from the JSON list    
                            df_Name_Matchs = pd.DataFrame(Name_Matchs)
                            df_Name_Matchs = df_Name_Matchs.drop(columns=['Trademark serial number','Trademark registration number','Trademark design phrase','reasoning'])  
                            # Create a table in the Word document    
                            table_Name_Matchs = document.add_table(df_Name_Matchs.shape[0] + 1, df_Name_Matchs.shape[1])
                            # Set a predefined table style (with borders)  
                            table_Name_Matchs.style = 'TableGrid'  # This is a built-in style that includes borders  
                            # Add the column names to the table    
                            for i, column_name in enumerate(df_Name_Matchs.columns):  
                                table_Name_Matchs.cell(0, i).text = column_name  
                            # Add the data to the table    
                            for i, row in df_Name_Matchs.iterrows():  
                                for j, value in enumerate(row):  
                                    table_Name_Matchs.cell(i + 1, j).text = str(value)

                if len(low_conflicts) > 0:  
                            document.add_heading('Trademarks with Low Conflicts:', level=2)  
                            # Create a pandas DataFrame from the JSON list    
                            df_low = pd.DataFrame(low_conflicts)  
                            df_low = df_low.drop(columns=['Trademark serial number','Trademark registration number','Trademark design phrase','reasoning'])
                            # Create a table in the Word document    
                            table_low = document.add_table(df_low.shape[0] + 1, df_low.shape[1])
                            # Set a predefined table style (with borders)  
                            table_low.style = 'TableGrid'  # This is a built-in style that includes borders  
                            # Add the column names to the table    
                            for i, column_name in enumerate(df_low.columns):  
                                table_low.cell(0, i).text = column_name  
                            # Add the data to the table    
                            for i, row in df_low.iterrows():  
                                for j, value in enumerate(row):  
                                    table_low.cell(i + 1, j).text = str(value)
                            
                def add_conflict_paragraph(document, conflict):  
                    p = document.add_paragraph(f"Trademark Name : {conflict.get('Trademark name', 'N/A')}")  
                    p.paragraph_format.line_spacing = Pt(18)  
                    p.paragraph_format.space_after = Pt(0)
                    p = document.add_paragraph(f"Trademark Status : {conflict.get('Trademark status', 'N/A')}")  
                    p.paragraph_format.line_spacing = Pt(18)  
                    p.paragraph_format.space_after = Pt(0)
                    p = document.add_paragraph(f"Trademark Owner : {conflict.get('Trademark owner', 'N/A')}")  
                    p.paragraph_format.line_spacing = Pt(18)  
                    p.paragraph_format.space_after = Pt(0)
                    p = document.add_paragraph(f"Trademark Class Number : {conflict.get('Trademark class Number', 'N/A')}")  
                    p.paragraph_format.line_spacing = Pt(18)
                    p.paragraph_format.space_after = Pt(0)  
                    p = document.add_paragraph(f"Trademark serial number : {conflict.get('Trademark serial number', 'N/A')}")  
                    p.paragraph_format.line_spacing = Pt(18)  
                    p.paragraph_format.space_after = Pt(0)
                    p = document.add_paragraph(f"Trademark registration number : {conflict.get('Trademark registration number', 'N/A')}")  
                    p.paragraph_format.line_spacing = Pt(18)  
                    p.paragraph_format.space_after = Pt(0) 
                    p = document.add_paragraph(f"Trademark Design phrase : {conflict.get('Trademark design phrase', 'N/A')}")  
                    p.paragraph_format.line_spacing = Pt(18)  
                    p.paragraph_format.space_after = Pt(0)
                    p = document.add_paragraph(" ")  
                    p.paragraph_format.line_spacing = Pt(18)  
                    p.paragraph_format.space_after = Pt(0) 
                    p = document.add_paragraph(f"{conflict.get('reasoning','N/A')}\n")  
                    p.paragraph_format.line_spacing = Pt(18)  
                    p = document.add_paragraph(" ")  
                    p.paragraph_format.line_spacing = Pt(18)  
                
                if len(high_conflicts) > 0:  
                    document.add_heading('Trademarks with High Conflicts Reasoning:', level=2)  
                    p = document.add_paragraph(" ")  
                    p.paragraph_format.line_spacing = Pt(18)  
                    for conflict in high_conflicts:  
                        add_conflict_paragraph(document, conflict)  
                
                if len(moderate_conflicts) > 0:  
                    document.add_heading('Trademarks with Moderate Conflicts Reasoning:', level=2)  
                    p = document.add_paragraph(" ")  
                    p.paragraph_format.line_spacing = Pt(18)  
                    for conflict in moderate_conflicts:  
                        add_conflict_paragraph(document, conflict)  
                        
                if len(Name_Matchs) > 0:  
                    document.add_heading("Trademarks with Name Match's Conflicts Reasoning:", level=2)  
                    p = document.add_paragraph(" ")  
                    p.paragraph_format.line_spacing = Pt(18)  
                    for conflict in Name_Matchs:  
                        add_conflict_paragraph(document, conflict)  
                
                if len(low_conflicts) > 0:  
                    document.add_heading('Trademarks with Low Conflicts Reasoning:', level=2)  
                    p = document.add_paragraph(" ")  
                    p.paragraph_format.line_spacing = Pt(18)  
                    for conflict in low_conflicts:  
                        add_conflict_paragraph(document, conflict)  
                        
                for i in range(70,96):
                    progress_bar.progress(i)  
                    
                progress_bar.progress(100)
    
                filename = proposed_name
                doc_stream = BytesIO()  
                document.save(doc_stream)  
                doc_stream.seek(0)  
                download_table = f'<a href="data:application/octet-stream;base64,{base64.b64encode(doc_stream.read()).decode()}" download="{filename + " Trademark Conflict Report"}.docx">Download: {filename}</a>'  
                st.sidebar.markdown(download_table, unsafe_allow_html=True)  
                st.success(f"{proposed_name} Document conflict report successfully completed!")
                
                end_time = time.time()
                elapsed_time = end_time - start_time
                elapsed_time = elapsed_time // 60 
                st.write(f"Time taken: {elapsed_time} mins")

                st.write("______________________________________________________________________________________________________________________________")
        
        progress_bar.progress(100)
        st.success("All documents processed successfully!")  
